"""Build a PLOS ONE submission-ready DOCX from docs/manuscript-draft.md.

Two things the source markdown intentionally keeps for repo/internal use but
that must not reach the journal: (1) the "Status" paragraph (draft history,
venue-selection rationale, links to internal gate-result docs), and (2) the
embedded figure images (PLOS requires figures as separate TIFF files, not
embedded in the manuscript document; docs/figures/Fig1.tif and Fig2.tif are
built by make_manuscript_figures.py). Both are stripped here, not from the
source .md, so the repo-facing draft stays readable on its own and the
submission file is generated, not hand-maintained separately.

Formatting applied after pandoc's markdown->docx conversion, since pandoc
alone does not set these Word-specific submission requirements: double
line spacing, continuous line numbering, and a page-number footer field.

Run: .venv/bin/python scripts/build_plos_docx.py
Writes: ../docs/ergofluids_manuscript_PLOS_ONE.docx
"""

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parents[1]
SRC_MD = REPO_ROOT.parent / "docs" / "manuscript-draft.md"
OUT_DOCX = REPO_ROOT.parent / "docs" / "ergofluids_manuscript_PLOS_ONE.docx"
TMP_MD = REPO_ROOT / "scripts" / "_render_docx" / "submission.md"


def strip_internal_content(text: str) -> str:
    # Drop the internal-tracking "Status" paragraph (draft history, venue
    # rationale, cross-references to internal gate-result docs).
    text = re.sub(
        r"\*\*Status:\*\*.*?those source documents\.\n",
        "",
        text,
        flags=re.DOTALL,
    )
    # Drop the two embedded figure images; PLOS wants figures as separate
    # files. The "Fig N." caption paragraphs immediately below each image
    # are kept, they are the actual captions PLOS wants in-text.
    text = re.sub(
        r"!\[.*?\]\(figures/fig\d_[a-z_]+\.png\)\n\n",
        "",
        text,
        flags=re.DOTALL,
    )
    return text


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_continuous_line_numbering(section) -> None:
    sect_pr = section._sectPr
    ln_num_type = OxmlElement("w:lnNumType")
    ln_num_type.set(qn("w:countBy"), "1")
    ln_num_type.set(qn("w:restart"), "continuous")
    ln_num_type.set(qn("w:distance"), "720")  # 0.5in, twentieths of a point
    sect_pr.append(ln_num_type)


def main() -> None:
    raw = SRC_MD.read_text()
    submission_md = strip_internal_content(raw)

    TMP_MD.parent.mkdir(parents=True, exist_ok=True)
    TMP_MD.write_text(submission_md)

    subprocess.run(
        [
            "pandoc",
            str(TMP_MD),
            "-o",
            str(OUT_DOCX),
            "--standalone",
            "--from",
            "markdown",
        ],
        check=True,
    )

    doc = Document(str(OUT_DOCX))

    # Double-space every paragraph (title page through references), the
    # PLOS ONE-required spacing for the whole manuscript body.
    for paragraph in doc.paragraphs:
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    for section in doc.sections:
        add_continuous_line_numbering(section)
        footer_paragraph = section.footer.paragraphs[0]
        footer_paragraph.alignment = 1  # center
        add_page_number_field(footer_paragraph)

    doc.save(str(OUT_DOCX))
    print(f"wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
